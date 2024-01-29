
# Author: Deepesh Patil
# Date: 04-02-2021
# Title: Library for file Encryption/Decryption
  

# IMPORTING REQUIRED LIBRARIES
import time
import emoji
import pwinput
import mysql.connector as mys
import random
import docx
import PyPDF2
import pyfiglet as fig
import openpyxl
import hashlib
from cryptography.fernet import Fernet
from colorama import Fore




# USING THE ASCII-ART TO DISPLAY THE APPLICATION NAME
def ascii_art():
    for i in fig.figlet_format(' Eagle - Encryptor ',font = 'slant',width = 200).split('\n'):
        print(Fore.LIGHTMAGENTA_EX + i.center(120) + Fore.GREEN)



# TO SET PASSWORD TO A PDF
def encrypt_pdf(file_path):
    """
    Encrypts a PDF file with a password and saves the result.

    Parameters:
    - file_path (str): Path to the input PDF file.
    - file_path (str): Path to save the encrypted PDF file.
    - password (str): Password for encryption.
    """

    print(Fore.GREEN + 'The PDF file will be locked with a Password.\n' + Fore.RESET)
    time.sleep(1)

    password = pwinput.pwinput(prompt ="Enter Password which You want to set to this PDF:\n", mask="*")

    with open(file_path, 'rb') as pdf_file:
        pdf_reader = PyPDF2.PdfReader(pdf_file)
        pdf_writer = PyPDF2.PdfWriter()

        for page_num in range(len(pdf_reader.pages)):
            pdf_writer.add_page(pdf_reader.pages[page_num])

        pdf_writer.encrypt(password)

        with open(file_path, 'wb') as result_pdf:
            pdf_writer.write(result_pdf)
    
    print(Fore.GREEN + 'Your PDF is now set with the above Password \n' + Fore.RESET)
    time.sleep(1)



# TO REMOVE THE PASSWORD SET ON THE PDF FILE
def decrypt_pdf(file_path):
    """
    Decrypts an encrypted PDF file with a password and saves the result.

    Parameters:
    - file_path (str): Path to the input encrypted PDF file.
    - file_path (str): Path to save the decrypted PDF file.
    - password (str): Password for decryption.
    """

    password = pwinput.pwinput(prompt ="Enter Password which You want to set to this PDF:\n", mask="*")
    try:
        with open(file_path, 'rb') as pdf_file:
            pdf_reader = PyPDF2.PdfReader(pdf_file)

            if pdf_reader.is_encrypted:
                pdf_reader.decrypt(password)
            else:
                print(Fore.RED + 'The File is not Encrypted.\n' + Fore.RESET)
                time.sleep(1)

            pdf_writer = PyPDF2.PdfWriter()

            for page_num in range(len(pdf_reader.pages)):
                pdf_writer.add_page(pdf_reader.pages[page_num])

            with open(file_path, 'wb') as result_pdf:
                pdf_writer.write(result_pdf)
        print(Fore.GREEN + 'The Password is Removed from the PDF.\n' + Fore.RESET)
        time.sleep(1)

    except:
        print(Fore.RED + 'Invalid Password !\nTry Again Later' + Fore.RESET)
        time.sleep(2)



# TO GET TEXT FROM A DOC FILE (TEXT-ONLY)
def getText_docx(file):
    doc = docx.Document(file)
    fullText = []
    for para in doc.paragraphs:
        fullText.append(para.text)
    return '\n'.join(fullText)



# TO CLEAR A DOC FILE
def clear_docx(file_path):

    # Creata an Empty Document
    doc = docx.Document()

    # Save the empty document to that path
    doc.save(file_path)



# ROT-13 MENTOD TO APPLY ON TEXT
def rot13(text):
    result = ''
    for char in text:
        if char.isalpha():
            shift = 13 if char.islower() else -13
            result += chr((ord(char) - ord('a' if char.islower() else 'A') + shift) % 26 + ord('a' if char.islower() else 'A'))
        elif char.isdigit():
            shift = 5
            result += str((int(char) + shift) % 10)
        else:
            result += char
    return result



# TO ADD A FILE TO YOUR LIST
def add_file():
    db = mys.connect(host = 'localhost', user = 'root', passwd = sql_pwd, database = 'files')
    cursor = db.cursor()

    cursor.execute('select fid from file_list')
    ids = cursor.fetchall()
    fid = random.randint(1,1000)
    while fid in ids:
        fid = random.randint(1,1000)
    fname = input('\nEnter the File Name :\n')
    p = input('Enter the Full File Path :\n')
    p = p.split('\\')
    fpath = ''
    for i in p:
        if i != p[-1]:
            fpath += i
            fpath += '\\\\'
        else:
            fpath += i
            break

    cursor.execute('insert into file_list values({0},"{1}","{2}")'.format(fid,fname,fpath))
    db.commit()

    print(Fore.GREEN + '\nFile Added to the List\n' + Fore.RESET)

    cursor.execute('select fpath from file_list where fname like "%{0}%"'.format(fname))
    return cursor.fetchall()



# TO VERIFY WHICH FILE IS TO BE ENCRYPTED/DECRYPTED
def which_file():
    try:
        db = mys.connect(host = 'localhost', user = 'root', passwd = sql_pwd, database = 'files')
        cursor = db.cursor()

        cursor.execute('select * from file_list')
        data = cursor.fetchall()

        if data == []:
            print(Fore.RED + 'No Files Added to the List.\n' + Fore.RESET)
            return None

        else:
            name = input('Enter the File Name or a Part of It :\n')
            cursor.execute('select * from file_list where fname like "%{0}%"'.format(name))
            data = cursor.fetchall()

            if data == []:
                print(Fore.RED + 'No Such File Found in the List.' + Fore.RESET)
                time.sleep(1)
                print('\nAdd the file to your list and try again\n')


            elif len(data) > 1:
                print('More than File Found!\n')
                cursor.execute('select * from file_list where fname like "%{0}%"'.format(name))
                data = cursor.fetchall()
                print('FILE ID','FILE NAME','FILE PATH',sep = '\t')
                print('-'*50)
                for a,b,c in data:
                    print(a,b,c,sep = '\t')
                print('')
                time.sleep(1)
                n = int(input('Enter File ID :\n'))
                cursor.execute('select * from file_list where fid = {0}'.format(n))
                f = cursor.fetchall()
                if f == []:
                    print(Fore.RED + 'File ID Not Found !\n' + Fore.RESET)
                else:
                    for a,b,c in f:
                        return c
            
            else:
                for a,b,c in data:
                    return c
    except:
        print(Fore.RED + 'An error occured while performing thiss operation.\nPlease contact the publisher.' + Fore.RESET)
        


# FUNCTION TO ENCRYPT A FILE
def encrypt():
    file_path = which_file()
    if file_path != None:
        try:
            ext = ''
            for i in file_path:
                if i == '.':
                    for j in range(file_path.index('.'),len(file_path)):
                        ext += file_path[j]

            if (ext == '.docx' or ext == '.doc'):
                while True:
                    if (input(Fore.RED + '\033[1mPlease Note, after encrypting a \'.docx\' format everyting other than the text in the document is deleted.\nMake sure that the document contains only text, and nothing else.\033[0m\n\nPress Enter to Continue or Press \'N\' to not proceed further : ' + Fore.RESET).lower() == 'n'):
                        break
                    
                    else:        
                        ddata = getText_docx(file_path)
                        clear_docx(file_path)
                        doc = docx.Document(file_path)
                        edata = rot13(ddata)
                        doc.add_paragraph(edata)
                        print(Fore.GREEN + 'File Encrypted !' + Fore.RESET)
                        doc.save(file_path)
                        break
            
            elif (ext == '.pdf'):
                encrypt_pdf(file_path)

            elif (ext == '.mp3' or ext == '.m4a'):
                enc_audio(file_path, key)
            
            elif (ext == '.txt' or ext == '.py' or  ext == '.html' or ext == '.htm' or ext == '.c' or ext == '.cpp' or ext == '.css' or ext == '.php'):
                f = open(file_path,'r')
                data = f.read()
                f.close()
                    
                f = open(file_path,'w')
                for i in data:
                    f.write(chr(ord(i)-5))
                f.close()
                print(Fore.GREEN + 'File Encrypted\n' + Fore.RESET)
            
            elif (ext == '.jpg' or ext == '.png' or ext == '.ico' or ext == '.jpeg' or ext == '.img'):
                enc_img(file_path)
            
            elif (ext == '.xlsx' or ext == '.xls' or ext == '.csv'):
                encrypt_excel_file(file_path, key)
            
            else:
                print('This Application does not support',ext,'type of files.\nTry some other format.\n')
                time.sleep(2)
        except:
            print(Fore.RED + 'An error occured while encrypting the file.\nTry verifying the file details/path and try again.' + Fore.RESET)
            time.sleep(2)
        


# FUNCTION TO DECRYPT A FILE
def decrypt():
    file_path = which_file()
    if file_path != None:
        try:
            ext = ''
            for i in file_path:
                if i == '.':
                    for j in range(file_path.index('.'),len(file_path)):
                        ext += file_path[j]

            if(ext == '.docx' or ext == '.doc'):
                edata = getText_docx(file_path)
                clear_docx(file_path)
                doc = docx.Document(file_path)
                ddata = rot13(edata)
                doc.add_paragraph(ddata)
                print(Fore.GREEN + 'File Decrypted !' + Fore.RESET)
                doc.save(file_path)
            
            elif (ext == '.pdf'):
                decrypt_pdf(file_path)
            
            elif (ext == '.mp3' or ext == '.m4a'):
                dec_audio(file_path, key)

            elif (ext == '.txt' or ext == '.py' or  ext == '.html' or ext == '.htm' or ext == '.c' or ext == '.cpp' or ext == '.css' or ext == '.php'):
                f = open(file_path,'r')
                data = f.read()
                f.close()

                f = open(file_path,'w')
                for i in data:
                    f.write(chr(ord(i)+5))
                f.close()
                print(Fore.GREEN + 'File Decrypted\n' + Fore.RESET)
            
            elif (ext == '.jpg' or ext == '.png' or ext == '.ico' or ext == '.jpeg' or ext == '.img'):
                dec_img(file_path)
            
            elif (ext == '.xlsx' or ext == '.xls' or ext == '.csv'):
                decrypt_excel_file(file_path, key)
            
            else:
                print('This Application does not support',ext,'type of files.\nTry some other format.\n')
                time.sleep(2)
        except:
            print(Fore.RED + 'An error occured while decrypting the file.\nTry verifying the file details/path and try again.' + Fore.RESET)
            time.sleep(2)



# FUNCTION THAT SHOWS ALL THE FILES PRESENT IN THE LIST
def view_file_list():
    try:   
        db = mys.connect(host = 'localhost', user = 'root', passwd = sql_pwd, database = 'files')
        cursor = db.cursor()
        cursor.execute('select * from file_list')
        data = cursor.fetchall()

        if data != []:        
            print('-'*6,'FILE NAME','-'*6,'\n')
            for a,b,c in data:
                print(b)
            print('\n')
        else:
            print(Fore.RED + 'No Files Added to the List.\n' + Fore.RESET)
    except:
        print(Fore.RED + 'Unable to connect to the MySQL Database.\nPlease verify the MySQL password, by running the setup file and try again later.' + Fore.RESET)



# FUNCTION TO ENCRYPT THE PASSWORD FILES
def encrypt1(path):
    f = open(path,'r')
    data = f.read()
    f.close()

    f = open(path,'w')
    for i in data:
        f.write(chr(ord(i)-5))
    f.close()



# FUNCTION TO DECRYPT THE PASSWORD FILES
def decrypt1(path):
    f = open(path,'r')
    data = f.read()
    f.close()

    f = open(path,'w')
    for i in data:
        f.write(chr(ord(i)+5))
    f.close()
 


# FUNCTION TO CHECK THE PASSWORD SET BY THE USER
def check():
    decrypt1('./!#%/pwd.txt')
    f = open('./!#%/pwd.txt','r')
    data = f.read()
    f.close()
    encrypt1('./!#%/pwd.txt')

    p = pwinput.pwinput(prompt ="\nENTER PASSWORD :\n", mask="*")
    print()
    if hashlib.sha256(p.encode()).hexdigest() == data:
        pass
    else:
        print(Fore.RED + 'INVALID PASSWORD !' + Fore.RESET) 
        print('   ',(emoji.emojize(":pensive_face:"))*3)
        time.sleep(5)
        quit()

             
   
# FUNCTION TO CHANGE THE USER SET PASSOWRD
def cng_pwd():
    decrypt1('./!#%/pwd.txt')
    f = open('./!#%/pwd.txt','r')
    data = f.read()
    f.close()
    encrypt1('./!#%/pwd.txt')

    o = pwinput.pwinput(prompt ="ENTER PREVIOUSLY SET PASSWORD :\n", mask="*")
    print()

    if hashlib.sha256(o.encode()).hexdigest() == data:
        decrypt1('./!#%/pwd.txt')
        f = open('./!#%/pwd.txt','w')
        p = pwinput.pwinput(prompt ="ENTER THE NEW PASSWORD :\n", mask="*")
        print()
        f.write(hashlib.sha256(p.encode()).hexdigest())
        print(Fore.GREEN + 'Password was Successfully Changed.\n' + Fore.RESET)
        f.close()
        encrypt1('./!#%/pwd.txt')
    else:
        print(Fore.RED + 'INVALID PASSWORD !' + Fore.RESET)
        print('   ',(emoji.emojize(":pensive_face:"))*3,'\n')
        pass
    time.sleep(2)



# FUNCTION WHICH DELETES A PARTICULAR FILE FROM THE LOCAL LIST
def delete_file():
    db = mys.connect(host = 'localhost', user = 'root', passwd = sql_pwd, database = 'files')
    cursor = db.cursor()
    cursor.execute('select fname from file_list')
    data = cursor.fetchall()

    try:
        if data == []:
            print(Fore.RED + 'No Files Added to the List.\n' + Fore.RESET)

        elif len(data) > 1:
            print('More than File Found!\n')

            cursor.execute('select * from file_list')
            data = cursor.fetchall()

            print('FILE ID','FILE NAME','FILE PATH',sep = '\t')
            print('-'*50)
            for a,b,c in data:
                print(a,b,c,sep = '\t')
            print('')
            time.sleep(1)

            n = int(input('Enter File ID :\n'))
            cursor.execute('select * from file_list where fid = {0}'.format(n))
            data = cursor.fetchall()
            if data != []:
                cursor.execute('delete from file_list where fid = {0}'.format(n))
                db.commit()
                print(Fore.GREEN + 'File Deleted\n' + Fore.RESET)
            else:
                print(Fore.RED + 'No File Present with the Given ID!\n' + Fore.RESET)

        else:
            cursor.execute('delete from file_list')
            db.commit()
            print(Fore.GREEN + 'File Deleted\n' + Fore.RESET)
    except:
        print(Fore.RED + 'An error occured while deleting the file\nTry Again Later !' + Fore.RESET)



# AN OUTRO FOR UI
def outro(path):
    f = open(path,'r')
    tag = f.readlines()
    print("\033[93m","\033[1m",tag[random.randint(0,len(tag))].upper(),"\033[0m")
    f.close()



# FUNCTION WHICH DELETES ALL THE FILES FROM THE LOCAL LIST
def delete_all():
    decrypt1('./!#%/pwd.txt')
    f = open('./!#%/pwd.txt','r')
    data = f.read()
    f.close()
    encrypt1('./!#%/pwd.txt')

    p = pwinput.pwinput(prompt ="\nPLEASE ENTER PASSWORD BEFORE PROCEEDING :\n", mask="*")
    print()
    while p == data:
        db = mys.connect(host = 'localhost', user = 'root', passwd = sql_pwd, database = 'files')
        cursor = db.cursor()

        cursor.execute('truncate file_list')
        print(Fore.GREEN + 'All Files Deleted !\n' + Fore.RESET)
        break

    else:
        print(Fore.RED + +'Invalid Password !\n' + Fore.RESET)
        print('   ',(emoji.emojize(":pensive_face:"))*3)



# AN INTRO FOR UI
def intro():
    print(Fore.LIGHTBLUE_EX + "\n\033[1m"," "*15,"WELCOME TO THE EAGLE-ENCRYPTOR")
    time.sleep(1)
    print("\033[91m\nPlease Make sure that all the files that you will be working on are CLOSED.\033[0m" + Fore.RESET)
    time.sleep(2)



# THE MENU FUNCTION
def menu():
    print(Fore.YELLOW + '\n\n\n','~'*16,' COMMAND LIST ','~'*16,'\n' + Fore.RESET)
    print(Fore.LIGHTMAGENTA_EX +'Commands\tTask/Uses' + Fore.RESET)
    print('-'*45)
    print('ls\t\tShows List of all Commands\nencf (ef)\tEncrypt a File from Your List\ndecf (df)\tDecrypt a File from Your List\nview (v)\tView List of Files \nadd (a)\t\tAdd a File in your List\ndel (d)\t\tDelete a File from Your List\ndelall (da)\tDelete All Files from the List\npwd (cp)\tChange Password\nquit (q)\tExit the Application\n')



# FUNCTION TO SETUP CONNECTION
def sql_setup(sql_pwd):
    try:
        decrypt1('./!#%/sql_pwd.txt')
        f = open('./!#%/sql_pwd.txt','w')
        f.write(hashlib.sha256(sql_pwd.encode()).hexdigest())
        f.close()
        encrypt1('./!#%/sql_pwd.txt')
        
        try:
            print('\nSQL Password Saved')
            time.sleep(0.5)
            db = mys.connect(host = 'localhost', user = 'root', passwd = sql_pwd)
            cursor = db.cursor()
            cursor.execute('CREATE DATABASE IF NOT EXISTS files')
            db.commit()
            print('Database Created')
            time.sleep(0.5)


            cursor.execute('USE files')
            cursor.execute('CREATE TABLE IF NOT EXISTS files (fid int primary key not null, fname varchar(100) not null, fpath varchar(200) not null)')
            db.commit()
            print('Table Created')

            cursor.execute('truncate files')
            db.commit()
            time.sleep(1)
        
        except:
            print(Fore.RED + 'An Unexpected error occured while creating the Database.\nPlease make sure that you have installed MySQL along with the Python-Connector (Steps mentioned in the README file).' + Fore.RESET)
        
    except:
        print(Fore.RED + 'An Unexpected Error Occured, please refer to the README file for manual installation.\n' + Fore.RESET)



# FUNCTION TO ENCRYPT AN IMAGE
def enc_img(fpath):
    key = 127
    try:
        # take path of image as a input
        # path = input(r'Enter path of Image : ')
        
        # taking encryption key as input
        # key = int(input('Enter Key for encryption of Image : '))
        
        # print path of image file and encryption key that
        # we are using
        # print('The path of file : ', fpath)
        # print('Key for encryption : ', key)
        
        # open file for reading purpose
        fin = open(fpath, 'rb')
        
        # storing image data in variable "image"
        image = fin.read()
        fin.close()
        
        # converting image into byte array to 
        # perform encryption easily on numeric data
        image = bytearray(image)

        # performing XOR operation on each value of bytearray
        for index, values in enumerate(image):
            image[index] = values ^ key

        # opening file for writing purpose
        fin = open(fpath, 'wb')
        
        # writing encrypted data in image
        fin.write(image)
        fin.close()
        time.sleep(0.5)
        print('Encryption Done...')

        
    except Exception:
        print(Fore.RED + 'Error caught : ', Exception.__name__ + Fore.RESET)



# FUNCTION TO DECRYPT AN IMAGE
def dec_img(fpath):
    key = 127
    try:
        # take path of image as a input
        # path = input(r'Enter path of Image : ')
        
        # taking decryption key as input
        # key = int(input('Enter Key for encryption of Image : '))
        
        # print path of image file and decryption key that we are using
        # print('The path of file : ', fpath)
        # print('Note : Encryption key and Decryption key must be same.')
        # print('Key for Decryption : ', key)
        
        # open file for reading purpose
        fin = open(fpath, 'rb')
        
        # storing image data in variable "image"
        image = fin.read()
        fin.close()
        
        # converting image into byte array to perform decryption easily on numeric data
        image = bytearray(image)

        # performing XOR operation on each value of bytearray
        for index, values in enumerate(image):
            image[index] = values ^ key

        # opening file for writing purpose
        fin = open(fpath, 'wb')
        
        # writing decryption data in image
        fin.write(image)
        fin.close()
        time.sleep(0.5)
        print('Decryption Done...')


    except Exception:
        print(Fore.RED + 'Error caught : ', Exception.__name__ + Fore.RESET)



# A Function to encrypt the text in Excel
def encrypt_excel_file(file_path, key):
    try:
        # Read Excel file into a DataFrame
        workbook = openpyxl.load_workbook(file_path)
        sheet = workbook.active
        data = [row for row in sheet.iter_rows(values_only=True)]

        # Convert data to a string for encryption
        data_str = str(data)

        # Encrypt data using Fernet symmetric key cryptography
        cipher_suite = Fernet(key)
        encrypted_data = cipher_suite.encrypt(data_str.encode())

        # Save the encrypted data to a new Excel file
        new_workbook = openpyxl.Workbook()
        new_sheet = new_workbook.active

        # Write the encrypted data to the new Excel file
        new_sheet.append(["Encrypted Data"])
        new_sheet.append([encrypted_data.decode()])

        # Save the new Excel file
        new_workbook.save(file_path)
        print(Fore.GREEN + 'Excel File Encrypted!' + Fore.RESET)
        time.sleep(2)
    except:
        print(Fore.RED + 'An unexpected error occured while performing the encryption.\nPlease try again after some time.' + Fore.RESET)



# A Function to decrypt the dext in excel
def decrypt_excel_file(file_path, key):
    try:
        # Load the encrypted Excel file
        workbook = openpyxl.load_workbook(file_path)
        sheet = workbook.active

        # Extract the encrypted data
        encrypted_data = sheet['A2'].value

        # Decrypt the data using Fernet symmetric key cryptography
        cipher_suite = Fernet(key)
        decrypted_data = cipher_suite.decrypt(encrypted_data.encode())

        # Convert the decrypted data string back to a list
        data = eval(decrypted_data.decode())

        # Save the decrypted data to a new Excel file
        new_workbook = openpyxl.Workbook()
        new_sheet = new_workbook.active
        for row in data:
            new_sheet.append(row)

        new_workbook.save(file_path)
        print(Fore.GREEN + 'Excel File Decrypted!' + Fore.RESET)
    except:
        print(Fore.RED + 'An unexpected error occured while performing the encryption.\nPlease try again after some time.' + Fore.RESET)



# A function to encrypt an audio file
def enc_audio(file_path, key):
    with open(file_path, 'rb') as f:
        data = f.read()
        f.close()
    
    with open(file_path,'wb') as f:
        f.write(Fernet(key).encrypt(data))
        f.close()
    
    print(Fore.GREEN + 'File Encrypted !' + Fore.RESET)
    time.sleep(1)



# A Function to decrypt an audio file
def dec_audio(file_path, key):
    with open(file_path, 'rb') as f:
        data = f.read()
        f.close()
    
    with open(file_path,'wb') as f:
        f.write(Fernet(key).decrypt(data))
        f.close()
    
    print(Fore.GREEN + 'File Decrypted !' + Fore.RESET)
    time.sleep(1)


 
# GETTING THE MYSQL PASSWORD and the KEY
flag = True

decrypt1('./!#%/key.txt')
f = open('./!#%/key.txt','r')
key = f.read()
f.close()
encrypt1('./!#%/key.txt')

try:
    f = open('./!#%/sql_pwd.txt','r')
    sql_pwd = f.read()
    f.close()
except:
    sql_pwd = ''
    
    
